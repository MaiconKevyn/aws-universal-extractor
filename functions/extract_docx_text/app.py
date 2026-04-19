import io
from typing import Any

from docx import Document

from app_common.exceptions import DocumentExtractionError
from app_common.logging import get_logger, log_json
from app_common.s3_utils import get_object_bytes, put_text


logger = get_logger(__name__)


def _paragraph_text(document) -> list[str]:
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def _table_to_text(table) -> str:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if not any(cells):
            continue
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def lambda_handler(event: dict[str, Any], context: Any) -> dict[str, Any]:
    bucket = event["document"]["bucket"]
    key = event["document"]["key"]
    output_bucket = event["artifacts"]["output_bucket"]
    output_prefix = event["artifacts"]["output_prefix"]

    docx_bytes = get_object_bytes(bucket, key)
    document = Document(io.BytesIO(docx_bytes))

    sections: list[str] = []
    paragraphs = _paragraph_text(document)
    if paragraphs:
        sections.append("=== Paragraphs ===\n" + "\n".join(paragraphs))

    table_summary: list[dict[str, Any]] = []
    for index, table in enumerate(document.tables, start=1):
        table_text = _table_to_text(table)
        if not table_text.strip():
            continue
        sections.append(f"=== Table {index} ===\n{table_text}")
        table_summary.append({
            "index": index,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
        })

    document_text = "\n\n".join(sections)

    if not document_text.strip():
        raise DocumentExtractionError("DOCX document is empty - no text or tables contain data")

    raw_text_key = f"{output_prefix}/raw_text.txt"
    put_text(output_bucket, raw_text_key, document_text)

    event["artifacts"]["raw_text"] = {
        "bucket": output_bucket,
        "key": raw_text_key,
    }
    event["docx_extraction"] = {
        "engine": "python-docx",
        "paragraph_count": len(paragraphs),
        "tables": table_summary,
        "text_length": len(document_text),
    }

    log_json(
        logger,
        "DOCX text extracted",
        request_id=event["request_id"],
        paragraph_count=len(paragraphs),
        table_count=len(table_summary),
        text_length=len(document_text),
    )
    return event
