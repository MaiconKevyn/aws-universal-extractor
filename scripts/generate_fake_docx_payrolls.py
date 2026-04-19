"""Generate synthetic US payroll DOCX files for testing the extraction pipeline.

Dev/runtime dependency:
    pip install python-docx faker

Usage:
    python scripts/generate_fake_docx_payrolls.py --out tests/fixtures/payroll/docx --count 5

Each DOCX is paired with a <stem>.expected.json ground-truth file using the same
schema as the PDF, XLSX, and CSV fixtures.
"""

from __future__ import annotations

import argparse
import json
import random
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from faker import Faker

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _paystub_data import Paystub, build_paystub, paystub_to_ground_truth  # noqa: E402


def _fmt_usd(value: float | None) -> str:
    if value is None:
        return ""
    return f"{value:,.2f}"


def _set_table_headers(table, headers: list[str]) -> None:
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def render_docx(p: Paystub, out_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading(p.employer_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(p.employer_address)
    doc.add_paragraph(f"EIN: {p.employer_ein}")
    doc.add_paragraph(f"EARNINGS STATEMENT    Pay Date: {p.pay_date}")

    info = doc.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    info_rows = [
        ("Employee", p.employee_name, "Employee ID", p.employee_id),
        ("Job Title", p.job_title, "SSN", f"XXX-XX-{p.ssn_last4}"),
        ("Hire Date", p.hire_date, "Pay Frequency", p.pay_frequency.title()),
        ("Pay Period", f"{p.period_start} to {p.period_end}", "Pay Rate", _fmt_usd(p.pay_rate)),
    ]
    for row, values in zip(info.rows, info_rows, strict=True):
        for index, value in enumerate(values):
            row.cells[index].text = value

    doc.add_heading("Earnings", level=2)
    earnings = [item for item in p.line_items if item.kind == "earning"]
    earn_table = doc.add_table(rows=1, cols=6)
    earn_table.style = "Table Grid"
    _set_table_headers(earn_table, ["Code", "Description", "Hours", "Rate", "Current", "YTD"])
    for item in earnings:
        cells = earn_table.add_row().cells
        values = [
            item.code or "",
            item.description,
            _fmt_usd(item.hours),
            _fmt_usd(item.rate),
            _fmt_usd(item.amount),
            _fmt_usd(item.ytd_amount),
        ]
        for index, value in enumerate(values):
            cells[index].text = value
    cells = earn_table.add_row().cells
    for index, value in enumerate(["", "Gross Pay", "", "", _fmt_usd(p.gross_pay), _fmt_usd(p.ytd_gross_pay)]):
        cells[index].text = value

    doc.add_heading("Deductions", level=2)
    deductions = [item for item in p.line_items if item.kind == "deduction"]
    ded_table = doc.add_table(rows=1, cols=4)
    ded_table.style = "Table Grid"
    _set_table_headers(ded_table, ["Code", "Description", "Current", "YTD"])
    for item in deductions:
        cells = ded_table.add_row().cells
        values = [item.code or "", item.description, _fmt_usd(item.amount), _fmt_usd(item.ytd_amount)]
        for index, value in enumerate(values):
            cells[index].text = value
    cells = ded_table.add_row().cells
    for index, value in enumerate([
        "",
        "Total Deductions",
        _fmt_usd(p.total_deductions),
        _fmt_usd(p.ytd_gross_pay - p.ytd_net_pay),
    ]):
        cells[index].text = value

    doc.add_heading("Summary", level=2)
    summary = doc.add_table(rows=3, cols=4)
    summary.style = "Table Grid"
    summary_rows = [
        ("Gross Pay", _fmt_usd(p.gross_pay), "YTD Gross", _fmt_usd(p.ytd_gross_pay)),
        ("Total Taxes", _fmt_usd(p.total_taxes), "Total Deductions", _fmt_usd(p.total_deductions)),
        ("", "", "NET PAY", _fmt_usd(p.net_pay)),
    ]
    for row, values in zip(summary.rows, summary_rows, strict=True):
        for index, value in enumerate(values):
            row.cells[index].text = value

    doc.save(out_path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--count", type=int, default=5)
    ap.add_argument("--seed", type=int, default=42)
    ap.add_argument("--locale", default="en_US")
    ap.add_argument(
        "--variant",
        choices=("canonical", "with_overtime", "with_bonus", "mixed"),
        default="mixed",
    )
    args = ap.parse_args()

    random.seed(args.seed)
    fake = Faker(args.locale)
    Faker.seed(args.seed)

    args.out.mkdir(parents=True, exist_ok=True)

    variants = (
        ["canonical", "with_overtime", "with_bonus"]
        if args.variant == "mixed"
        else [args.variant]
    )

    for i in range(args.count):
        variant = variants[i % len(variants)]
        p = build_paystub(fake, variant=variant)
        stem = f"paystub_{i + 1:03d}_{variant}"
        docx_path = args.out / f"{stem}.docx"
        gt_path = args.out / f"{stem}.expected.json"

        render_docx(p, docx_path)
        gt_path.write_text(
            json.dumps(paystub_to_ground_truth(p), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"wrote {docx_path}")


if __name__ == "__main__":
    main()
