"""Smoke-test CSV and DOCX extraction end-to-end against local fixtures.

Mirrors the Lambda text-extraction logic (no S3), then calls OpenAI with the
payroll/v1 profile and compares to the ground-truth JSON. Reports field-level
accuracy so we can confirm the schema is truly format-agnostic.

Usage:
    ./.venv/bin/python scripts/smoke_test_formats.py
"""

from __future__ import annotations

import csv as csv_module
import io
import json
import math
import sys
from pathlib import Path

from docx import Document
from jsonschema import Draft202012Validator

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "layers" / "common" / "python"))

from app_common.openai_client import OpenAIExtractionClient  # noqa: E402
from app_common.profiles import load_profile  # noqa: E402
from app_common.validators import to_metadata_json  # noqa: E402


def extract_csv_text(path: Path) -> str:
    data = path.read_bytes()
    text = data.decode("utf-8-sig")
    try:
        dialect = csv_module.Sniffer().sniff(text[:8192])
    except csv_module.Error:
        dialect = csv_module.excel
    rows = [[(c or "").strip() for c in row] for row in csv_module.reader(io.StringIO(text), dialect)]
    rows = [r for r in rows if any(r)]
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows)
    return f"=== CSV: {path.name} ===\n{body}"


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    sections: list[str] = []
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        sections.append("=== Paragraphs ===\n" + "\n".join(paragraphs))
    for i, table in enumerate(doc.tables, start=1):
        lines = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
        if lines:
            sections.append(f"=== Table {i} ===\n" + "\n".join(lines))
    return "\n\n".join(sections)


def flatten(obj, prefix=""):
    out = {}
    if isinstance(obj, dict):
        for k, v in obj.items():
            out.update(flatten(v, f"{prefix}.{k}" if prefix else k))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.update(flatten(v, f"{prefix}[{i}]"))
    else:
        out[prefix] = obj
    return out


def values_match(a, b) -> bool:
    if a is None and b is None:
        return True
    if isinstance(a, (int, float)) and isinstance(b, (int, float)):
        return math.isclose(float(a), float(b), rel_tol=1e-3, abs_tol=0.01)
    return str(a).strip() == str(b).strip()


def accuracy(predicted: dict, expected: dict) -> tuple[float, list[str]]:
    p_flat = flatten(predicted)
    e_flat = flatten(expected)
    keys = sorted(set(p_flat) | set(e_flat))
    mismatches = []
    hits = 0
    for k in keys:
        if values_match(p_flat.get(k), e_flat.get(k)):
            hits += 1
        else:
            mismatches.append(f"  {k}: expected={e_flat.get(k)!r}  got={p_flat.get(k)!r}")
    return hits / len(keys) if keys else 1.0, mismatches


def run_case(label: str, text: str, expected_path: Path, profile: dict, client: OpenAIExtractionClient) -> None:
    print(f"\n=== {label} ===")
    print(f"  extracted {len(text)} chars")
    context = {
        "client_id": "smoke",
        "document_id": expected_path.stem.replace(".expected", ""),
        "metadata_json": to_metadata_json({}),
        "document_text": text,
    }
    result = client.extract(profile=profile, document_text=text, context=context)
    predicted = result["data"]
    expected = json.loads(expected_path.read_text(encoding="utf-8"))

    validator = Draft202012Validator(profile["schema"])
    schema_errors = list(validator.iter_errors(predicted))
    score, mismatches = accuracy(predicted, expected)

    print(f"  schema: {'OK' if not schema_errors else f'{len(schema_errors)} errors'}")
    print(f"  accuracy: {score * 100:.1f}%")
    if mismatches:
        print(f"  mismatches ({len(mismatches)}):")
        for m in mismatches[:15]:
            print(m)


def main() -> None:
    profile = load_profile(str(REPO_ROOT / "profiles"), "payroll", "v1")
    client = OpenAIExtractionClient()

    csv_path = REPO_ROOT / "tests/fixtures/payroll/csv/paystub_001_canonical.csv"
    docx_path = REPO_ROOT / "tests/fixtures/payroll/docx/paystub_001_canonical.docx"

    run_case(
        "CSV paystub_001_canonical",
        extract_csv_text(csv_path),
        csv_path.with_suffix(".expected.json").with_name(csv_path.stem + ".expected.json"),
        profile,
        client,
    )
    run_case(
        "DOCX paystub_001_canonical",
        extract_docx_text(docx_path),
        docx_path.with_name(docx_path.stem + ".expected.json"),
        profile,
        client,
    )


if __name__ == "__main__":
    main()
